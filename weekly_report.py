"""
US Residential Intelligence v2 — weekly_report.py
articles.csv 기반 인텔리전스 리포트(주간/월간/분기/반기) 자동 생성.

사용법:
  python weekly_report.py                  # 기본: 최근 7일 (weekly)
  python weekly_report.py --period monthly  # 최근 30일
  python weekly_report.py --period quarterly
  python weekly_report.py --period semi-annual
"""

import argparse
import csv
import os
import re
import sys
from datetime import datetime, timedelta, timezone

import anthropic
from dotenv import load_dotenv

load_dotenv()

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ------------------------------------------------------------------
# 설정
# ------------------------------------------------------------------

ARTICLES_CSV = "articles.csv"
REPORTS_DIR  = "reports"
MODEL        = "claude-sonnet-4-5"
MAX_TOKENS   = 4000

PERIOD_DAYS = {
    "weekly":       7,
    "monthly":      30,
    "quarterly":    90,
    "semi-annual": 180,
}

PERIOD_LABEL = {
    "weekly":       "주간",
    "monthly":      "월간",
    "quarterly":    "분기",
    "semi-annual":  "반기",
}

SYSTEM_PROMPT = (
    "당신은 우미글로벌 해외사업팀의 미국 주거시장 리서치 애널리스트입니다. "
    "수집된 뉴스 기사를 바탕으로 경영진과 팀원이 읽을 원페이저 인텔리전스 리포트를 작성합니다."
)

REPORT_TEMPLATE = """\
아래는 {period_label} 수집된 미국 주거시장 뉴스 기사 {count}건입니다.
기간: {date_from} ~ {date_to}

{articles_text}

---

위 기사들을 바탕으로 아래 4개 섹션으로 구성된 한국어 인텔리전스 리포트를 작성해주세요.
각 섹션은 3~5개 bullet point로 핵심만 간결하게 정리합니다.
수치·규모·지역 정보가 있으면 반드시 포함하세요.

## 1. 개발 현황
섹터별 주요 개발 프로젝트, 지역, unit 수 (가능한 경우), 주목할 움직임.

## 2. 정책·이슈
인허가·세금·금리 관련 주요 정부 정책 및 시장 이슈.

## 3. 거래현황
상품별 거래 규모, 주요 자본 흐름 (domestic vs. 글로벌), 주요 플레이어.

## 4. 시사점
우미글로벌 해외사업팀 관점에서의 전략적 함의. BTR/Student Housing/GP 파트너십 기회 중심.
"""


# ------------------------------------------------------------------
# 데이터 로드 및 필터
# ------------------------------------------------------------------

def load_articles() -> list[dict]:
    if not os.path.exists(ARTICLES_CSV):
        return []
    with open(ARTICLES_CSV, encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def filter_by_period(articles: list[dict], days: int) -> list[dict]:
    cutoff = datetime.now(tz=timezone.utc) - timedelta(days=days)
    result = []
    for a in articles:
        raw = a.get("published_at", "")
        if not raw:
            continue
        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            if dt >= cutoff:
                result.append(a)
        except ValueError:
            continue
    return result


def format_articles(articles: list[dict]) -> str:
    lines = []
    for i, a in enumerate(articles, 1):
        tags = a.get("event_tags", "")
        lines.append(
            f"{i}. [{a.get('category','')} / {a.get('sector','')}] "
            f"{a.get('title','')}\n"
            f"   요약: {a.get('summary','')[:200]}\n"
            f"   태그: {tags}"
        )
    return "\n\n".join(lines)


# ------------------------------------------------------------------
# weekly .md 파일 로드 (분기/반기 재요약용)
# ------------------------------------------------------------------

def load_weekly_reports(days: int) -> list[tuple[str, str]]:
    """reports/ 폴더에서 기간 내 weekly .md 파일을 (파일명, 내용) 리스트로 반환."""
    if not os.path.isdir(REPORTS_DIR):
        return []
    cutoff = datetime.now() - timedelta(days=days)
    result = []
    for fname in sorted(os.listdir(REPORTS_DIR)):
        if not (fname.endswith("_weekly.md") or fname.endswith("-weekly.md")):
            continue
        # 파일명에서 날짜 추출 (YYYYMMDD_weekly.md)
        m = re.match(r"^(\d{8})_weekly\.md$", fname)
        if not m:
            continue
        try:
            fdate = datetime.strptime(m.group(1), "%Y%m%d")
        except ValueError:
            continue
        if fdate >= cutoff:
            fpath = os.path.join(REPORTS_DIR, fname)
            with open(fpath, encoding="utf-8") as f:
                result.append((fname, f.read()))
    return result


def generate_report_from_weeklies(weekly_reports: list[tuple[str, str]], period: str) -> str:
    """weekly .md 파일들을 종합해 분기/반기 리포트를 생성한다."""
    period_label = PERIOD_LABEL[period]
    weeks_count  = len(weekly_reports)

    combined = ""
    for fname, content in weekly_reports:
        combined += f"\n\n{'='*60}\n📄 {fname}\n{'='*60}\n{content}"

    user_prompt = (
        f"다음은 최근 {weeks_count}주간의 주간 리포트입니다.\n"
        f"이를 종합해 {period_label} 인텔리전스 리포트를 작성해주세요.\n\n"
        f"요청 사항:\n"
        f"- 4개 섹션(개발 현황 / 정책·이슈 / 거래현황 / 시사점)으로 구성\n"
        f"- 각 섹션은 3~7개 bullet point, 수치·지역·규모 정보 포함\n"
        f"- 반복 언급된 트렌드는 강조, 일회성 사건은 축약\n"
        f"- 시사점은 우미글로벌 해외사업팀 관점, BTR/Student Housing/GP 기회 중심\n\n"
        f"{combined}"
    )

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다.")

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )
    return response.content[0].text.strip()


# ------------------------------------------------------------------
# 리포트 생성
# ------------------------------------------------------------------

def generate_report(articles: list[dict], period: str) -> str:
    days = PERIOD_DAYS[period]

    date_to   = datetime.now().strftime("%Y-%m-%d")
    date_from = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")

    # woomi_relevance 높음 우선, 최대 100건으로 제한
    sorted_articles = sorted(
        articles,
        key=lambda a: (0 if a.get("woomi_relevance") == "높음" else 1 if a.get("woomi_relevance") == "보통" else 2)
    )[:100]
    articles_text = format_articles(sorted_articles)
    user_prompt = REPORT_TEMPLATE.format(
        period_label=PERIOD_LABEL[period],
        count=len(articles),
        date_from=date_from,
        date_to=date_to,
        articles_text=articles_text,
    )

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다.")

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )
    return response.content[0].text.strip()


# ------------------------------------------------------------------
# 저장
# ------------------------------------------------------------------

def save_report(content: str, period: str) -> str:
    os.makedirs(REPORTS_DIR, exist_ok=True)

    today    = datetime.now().strftime("%Y%m%d")
    filename = f"{today}_{period}.md"
    filepath = os.path.join(REPORTS_DIR, filename)

    days      = PERIOD_DAYS[period]
    date_from = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
    date_to   = datetime.now().strftime("%Y-%m-%d")
    generated = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    meta = (
        f"---\n"
        f"기간: {date_from} ~ {date_to} ({PERIOD_LABEL[period]})\n"
        f"생성: {generated}\n"
        f"---\n\n"
    )

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(meta + content)

    return filepath


# ------------------------------------------------------------------
# docx 변환
# ------------------------------------------------------------------

def _strip_md_bold(text: str) -> str:
    """**bold** 마크다운을 일반 텍스트로 변환 (run 분리 전 단순 제거)."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def save_report_docx(md_content: str, period: str) -> str:
    """마크다운 리포트를 우미글로벌 양식 .docx 로 저장한다."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx가 설치되지 않았습니다. pip install python-docx 를 실행하세요.")

    # ── 색상 상수 ──────────────────────────────────────────────
    C_NAVY   = RGBColor(0x1B, 0x3A, 0x5C)   # 진네이비
    C_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # 섹션 헤딩
    C_GREY   = RGBColor(0x99, 0x99, 0x99)   # 보조 텍스트
    C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    FONT     = "Noto Sans KR"

    days      = PERIOD_DAYS[period]
    date_from = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
    date_to   = datetime.now().strftime("%Y-%m-%d")
    year_month = datetime.now().strftime("%Y년 %-m월") if os.name != "nt" else datetime.now().strftime("%Y년 %#m월")

    # ── 헬퍼 ──────────────────────────────────────────────────
    def _run(para, text, size_pt=10, bold=False, color=None, italic=False):
        run = para.add_run(text)
        run.font.name  = FONT
        run.font.size  = Pt(size_pt)
        run.font.bold  = bold
        run.font.italic = italic
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
        if color:
            run.font.color.rgb = color
        return run

    def _para(doc, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        return p

    def _heading_border(para, color_hex="1B3A5C"):
        """단락 아래에 얇은 보더 추가."""
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _page_break(doc):
        p   = doc.add_paragraph()
        run = p.add_run()
        run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

    def _add_bold_runs(para, text, size_pt=10, base_color=None):
        """**bold** 구간을 run 분리해 추가한다."""
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                _run(para, part[2:-2], size_pt=size_pt, bold=True, color=base_color)
            else:
                _run(para, part, size_pt=size_pt, color=base_color)

    # ── 페이지 설정 (A4, 1인치 여백) ──────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── 헤더 ──────────────────────────────────────────────────
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.clear()
    _run(hp, "우미글로벌 해외사업팀 | 대외비", size_pt=10, color=C_GREY)
    _heading_border(hp)

    # ── 푸터 ──────────────────────────────────────────────────
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 페이지 번호 필드
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2  = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3  = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    fr = fp.add_run()
    for el in (fldChar1, instrText, fldChar2, fldChar3):
        fr._r.append(el)
    fr.font.size = Pt(9)
    fr.font.name = FONT
    # 작성 연월 우측
    fp.add_run("\t\t")
    _run(fp, year_month, size_pt=9, color=C_GREY)

    # ── 표지 ──────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_company = _para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_company, "우미글로벌 해외사업팀", size_pt=16, bold=True, color=C_NAVY)

    doc.add_paragraph()

    p_title = _para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_title, f"미국 주거시장 {PERIOD_LABEL[period]} 인텔리전스 리포트", size_pt=20, bold=True)

    doc.add_paragraph()

    p_meta = _para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_meta, f"기간: {date_from} ~ {date_to}", size_pt=10, color=C_GREY)

    doc.add_paragraph()

    p_date = _para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_date, f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}  |  해외사업팀", size_pt=10, color=C_GREY)

    # 표지 뒤 페이지 구분
    from docx.oxml import OxmlElement as OE
    p_br = doc.add_paragraph()
    run_br = p_br.add_run()
    br = OE("w:br")
    br.set(qn("w:type"), "page")
    run_br._r.append(br)

    # ── 본문 파싱 ──────────────────────────────────────────────
    lines = md_content.splitlines()
    in_frontmatter = False
    frontmatter_count = 0

    for line in lines:
        # YAML front-matter 스킵
        if line.strip() == "---":
            frontmatter_count += 1
            in_frontmatter = frontmatter_count < 2
            continue
        if in_frontmatter:
            continue

        # H1 (# ) — 문서 타이틀 줄, 표지와 중복되므로 얇게 처리
        if line.startswith("# "):
            p = _para(doc)
            _run(p, _strip_md_bold(line[2:].strip()), size_pt=14, bold=True, color=C_NAVY)
            _heading_border(p)

        # H2 (## ) — 섹션 제목
        elif line.startswith("## "):
            doc.add_paragraph()
            p = _para(doc)
            _run(p, _strip_md_bold(line[3:].strip()), size_pt=13, bold=True, color=C_BLUE)
            _heading_border(p, "2E75B6")

        # H3 (### )
        elif line.startswith("### "):
            p = _para(doc)
            _run(p, _strip_md_bold(line[4:].strip()), size_pt=12, bold=True, color=RGBColor(0x1F, 0x4D, 0x78))

        # Bullet (•, -, *)
        elif line.startswith("•") or (line.startswith("- ") and line.strip() != "---") \
                or line.startswith("* "):
            text = re.sub(r"^[•\-\*]\s*", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            _add_bold_runs(p, text, size_pt=10)

        # 들여쓰기 sub-bullet
        elif line.startswith("  - ") or line.startswith("  • "):
            text = re.sub(r"^\s+[•\-]\s*", "", line).strip()
            p = doc.add_paragraph(style="List Bullet 2")
            _add_bold_runs(p, text, size_pt=10)

        # **로만 구성된 줄 (강조 단락)
        elif line.strip().startswith("**") and line.strip().endswith("**") and line.strip() != "**":
            p = _para(doc)
            _run(p, line.strip()[2:-2], size_pt=10, bold=True)

        # 수평선 (---)
        elif line.strip() == "---":
            _heading_border(_para(doc))

        # 일반 텍스트
        elif line.strip():
            p = _para(doc)
            _add_bold_runs(p, line.strip(), size_pt=10)

    # ── 저장 ──────────────────────────────────────────────────
    os.makedirs(REPORTS_DIR, exist_ok=True)
    today    = datetime.now().strftime("%Y%m%d")
    filepath = os.path.join(REPORTS_DIR, f"{today}_{period}.docx")
    doc.save(filepath)
    return filepath


# ------------------------------------------------------------------
# 메인
# ------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="US Residential Intelligence 리포트 생성")
    parser.add_argument(
        "--period",
        choices=list(PERIOD_DAYS.keys()),
        default="weekly",
        help="리포트 기간 (기본값: weekly)",
    )
    parser.add_argument(
        "--export",
        choices=["docx"],
        default=None,
        help="추가 출력 포맷 (예: --export docx)",
    )
    args = parser.parse_args()

    period = args.period
    export = args.export
    days   = PERIOD_DAYS[period]

    print(f"=== US Residential Intelligence — {PERIOD_LABEL[period]} 리포트 생성 ===")
    print(f"기간: 최근 {days}일\n")

    # ── 분기/반기: weekly .md 재요약, fallback → articles.csv 직접 분석 ──
    use_weekly_summary = period in ("quarterly", "semi-annual")
    report_content = None

    if use_weekly_summary:
        weekly_reports = load_weekly_reports(days)
        if len(weekly_reports) < 3:
            print(
                f"[WARN] reports/ 내 weekly .md 파일이 {len(weekly_reports)}개로 부족합니다 "
                f"(최소 3개 필요). articles.csv 직접 분석으로 fallback합니다."
            )
            use_weekly_summary = False
        else:
            print(f"주간 리포트 {len(weekly_reports)}개를 종합해 {PERIOD_LABEL[period]} 리포트를 생성합니다.")
            print("Claude API 호출 중...\n")
            try:
                report_content = generate_report_from_weeklies(weekly_reports, period)
            except RuntimeError as e:
                print(f"오류: {e}")
                return

    if not use_weekly_summary:
        articles = load_articles()
        if not articles:
            print("articles.csv가 없거나 비어 있습니다. collector.py를 먼저 실행하세요.")
            return

        filtered = filter_by_period(articles, days)
        if not filtered:
            print(f"최근 {days}일 내 기사가 없습니다. 리포트를 생성하지 않습니다.")
            return

        print(f"대상 기사: {len(filtered)}건")
        print("Claude API 호출 중...\n")

        try:
            report_content = generate_report(filtered, period)
        except RuntimeError as e:
            print(f"오류: {e}")
            return

    filepath = save_report(report_content, period)
    print(f"리포트 저장 완료: {filepath}")

    if export == "docx":
        try:
            docx_path = save_report_docx(report_content, period)
            print(f"Word 문서 저장 완료: {docx_path}")
        except RuntimeError as e:
            print(f"[WARN] docx 변환 실패: {e}")

    print("\n--- 리포트 미리보기 (처음 500자) ---")
    print(report_content[:500])


if __name__ == "__main__":
    main()
